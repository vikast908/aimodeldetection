import io
import os
import re
import tempfile
import zipfile
from collections import Counter
from datetime import datetime, timedelta
from typing import Dict, List, Optional
import xml.etree.ElementTree as etree

import docx


WORD_RE = re.compile(r"\b\w+\b")


def count_words(text: str) -> int:
    return len(WORD_RE.findall(text or ""))


def parse_text_content(text: str, filename: str) -> Dict:
    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", normalized) if p.strip()]
    return {
        "filename": filename,
        "text": normalized,
        "paragraphs": paragraphs,
        "metadata": {},
        "track_changes": [],
        "font_info": {},
        "style_info": {},
        "file_type": filename.split(".")[-1].lower(),
    }


def parse_docx_bytes(data: bytes, filename: str) -> Dict:
    doc = docx.Document(io.BytesIO(data))
    paragraphs = [p.text for p in doc.paragraphs if p.text is not None]
    text = "\n".join(paragraphs).strip()

    font_info = extract_font_info(doc)
    style_info = extract_style_info(doc)
    metadata = extract_docx_metadata(data)
    track_changes = extract_track_changes(data)

    return {
        "filename": filename,
        "text": text,
        "paragraphs": paragraphs,
        "metadata": metadata,
        "track_changes": track_changes,
        "font_info": font_info,
        "style_info": style_info,
        "file_type": "docx",
    }


def parse_doc_bytes(data: bytes, filename: str) -> Dict:
    text = ""
    temp_path = None
    try:
        import textract

        with tempfile.NamedTemporaryFile(delete=False, suffix=".doc") as tmp:
            tmp.write(data)
            tmp.flush()
            temp_path = tmp.name
        extracted = textract.process(temp_path)
        if isinstance(extracted, bytes):
            text = extracted.decode("utf-8", errors="ignore")
        else:
            text = str(extracted)
    except Exception:
        try:
            text = data.decode("latin-1", errors="ignore")
        except Exception:
            text = ""
    finally:
        if temp_path:
            try:
                os.unlink(temp_path)
            except Exception:
                pass

    text = (text or "").strip()
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
    return {
        "filename": filename,
        "text": text,
        "paragraphs": paragraphs,
        "metadata": {},
        "track_changes": [],
        "font_info": {},
        "style_info": {},
        "file_type": "doc",
    }


def parse_document_bytes(data: bytes, filename: str) -> Dict:
    lower = filename.lower()
    if lower.endswith(".docx"):
        return parse_docx_bytes(data, filename)
    if lower.endswith(".doc"):
        return parse_doc_bytes(data, filename)
    if lower.endswith(".txt") or lower.endswith(".md"):
        text = data.decode("utf-8", errors="ignore")
        return parse_text_content(text, filename)
    if lower.endswith(".pdf"):
        return parse_pdf_bytes(data, filename)
    text = data.decode("utf-8", errors="ignore")
    return parse_text_content(text, filename)


def parse_pdf_bytes(data: bytes, filename: str) -> Dict:
    try:
        from PyPDF2 import PdfReader
    except Exception:
        text = data.decode("utf-8", errors="ignore")
        return parse_text_content(text, filename)

    reader = PdfReader(io.BytesIO(data))
    pages = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception:
            continue
    text = "\n\n".join(pages).strip()
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
    return {
        "filename": filename,
        "text": text,
        "paragraphs": paragraphs,
        "metadata": {},
        "track_changes": [],
        "font_info": {},
        "style_info": {},
        "file_type": "pdf",
    }


def extract_font_info(doc: docx.Document) -> Dict:
    run_fonts: List = []
    for para in doc.paragraphs:
        for run in para.runs:
            name = run.font.name
            size = run.font.size.pt if run.font.size else None
            if name or size:
                run_fonts.append((name or "unknown", size or 0))

    if not run_fonts:
        return {"clusters": 0, "dominant": None, "font_counts": {}}

    font_counts = Counter(run_fonts)
    dominant = font_counts.most_common(1)[0][0]
    clusters = 0
    in_cluster = False
    for font in run_fonts:
        if font != dominant:
            if not in_cluster:
                clusters += 1
                in_cluster = True
        else:
            in_cluster = False

    return {
        "clusters": clusters,
        "dominant": dominant,
        "font_counts": dict(font_counts),
    }


def extract_style_info(doc: docx.Document) -> Dict:
    heading_styles = set()
    list_styles = set()
    spacing_values = set()

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        if style_name.lower().startswith("heading"):
            heading_styles.add(style_name)

        if para._p is not None and para._p.pPr is not None and para._p.pPr.numPr is not None:
            num_id = para._p.pPr.numPr.numId
            if num_id is not None:
                list_styles.add(str(num_id.val))

        space_before = para.paragraph_format.space_before
        space_after = para.paragraph_format.space_after
        if space_before is not None:
            spacing_values.add(space_before.pt)
        if space_after is not None:
            spacing_values.add(space_after.pt)

    return {
        "heading_styles": sorted(heading_styles),
        "list_styles": sorted(list_styles),
        "spacing_values": sorted(spacing_values),
    }


def extract_docx_metadata(data: bytes) -> Dict:
    metadata: Dict = {}
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as zf:
            if "docProps/core.xml" in zf.namelist():
                core = zf.read("docProps/core.xml")
                metadata.update(parse_core_props(core))
            if "docProps/app.xml" in zf.namelist():
                app = zf.read("docProps/app.xml")
                metadata.update(parse_app_props(app))
    except Exception:
        return metadata
    return metadata


def parse_core_props(core_xml: bytes) -> Dict:
    ns = {
        "cp": "http://schemas.openxmlformats.org/package/2006/metadata/core-properties",
        "dcterms": "http://purl.org/dc/terms/",
    }
    root = etree.fromstring(core_xml)
    created = root.findtext(".//dcterms:created", namespaces=ns)
    modified = root.findtext(".//dcterms:modified", namespaces=ns)
    revision = root.findtext(".//cp:revision", namespaces=ns)
    metadata: Dict = {}
    if created:
        metadata["created"] = parse_iso_datetime(created)
    if modified:
        metadata["modified"] = parse_iso_datetime(modified)
    if revision and revision.isdigit():
        metadata["revision"] = int(revision)
    return metadata


def parse_app_props(app_xml: bytes) -> Dict:
    ns = {"ep": "http://schemas.openxmlformats.org/officeDocument/2006/extended-properties"}
    root = etree.fromstring(app_xml)
    total_time = root.findtext(".//ep:TotalTime", namespaces=ns)
    metadata: Dict = {}
    if total_time and total_time.isdigit():
        metadata["editing_minutes"] = int(total_time)
    return metadata


def parse_iso_datetime(value: str) -> Optional[datetime]:
    try:
        return datetime.fromisoformat(value.replace("Z", "+00:00"))
    except Exception:
        return None


def extract_track_changes(data: bytes) -> List[Dict]:
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as zf:
            if "word/document.xml" not in zf.namelist():
                return []
            xml_bytes = zf.read("word/document.xml")
    except Exception:
        return []

    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    try:
        root = etree.fromstring(xml_bytes)
    except Exception:
        return []

    # Register namespace for ElementTree
    for prefix, uri in ns.items():
        etree.register_namespace(prefix, uri)

    paragraphs = root.findall(".//{%s}p" % ns["w"])
    para_index = {id(p): idx for idx, p in enumerate(paragraphs)}

    edits: List[Dict] = []
    for tag in ("ins", "del"):
        for el in root.findall(".//{%s}%s" % (ns["w"], tag)):
            # Extract text from all text nodes
            text_nodes = []
            for t_elem in el.findall(".//{%s}t" % ns["w"]):
                if t_elem.text:
                    text_nodes.append(t_elem.text)
            for dt_elem in el.findall(".//{%s}delText" % ns["w"]):
                if dt_elem.text:
                    text_nodes.append(dt_elem.text)
            text = "".join(text_nodes)
            words = WORD_RE.findall(text)

            # Find parent paragraph
            paragraph_node = el
            while paragraph_node is not None and paragraph_node.tag != "{%s}p" % ns["w"]:
                parent = None
                for p in paragraphs:
                    if el in list(p.iter()):
                        parent = p
                        break
                paragraph_node = parent
                break

            para_idx = para_index.get(id(paragraph_node)) if paragraph_node is not None else None
            edits.append(
                {
                    "type": tag,
                    "word_count": len(words),
                    "paragraph_index": para_idx,
                    "text": text,
                }
            )
    return edits
